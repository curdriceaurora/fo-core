"""SafeDir rail — CI test (security hardening tracking: #264).

Companion to ``scripts/check_safedir_required.py``. The rail is in **advisory
mode** for PR1 — every test in this file passes regardless of how many
violations the detector finds. As PRs #267 / #268 / #269 / #270 migrate
call sites to SafeDir, the rail's enforcement scope expands per-directory.

The CI test verifies three things:

1. The detector itself works on synthetic inputs (positive + negative cases).
2. Running ``check_safedir_required.py --advisory`` against the live tree
   exits 0 (no false-positive on the implementation modules).
3. The current baseline count is recorded so any regression that adds new
   call sites in PRs targeting unrelated areas is visible in CI logs.
"""

from __future__ import annotations

import ast
import subprocess
import sys
from pathlib import Path

import pytest

pytestmark = pytest.mark.ci

_FO_ROOT = Path(__file__).resolve().parents[2]
_SCRIPT = _FO_ROOT / "scripts" / "check_safedir_required.py"

sys.path.insert(0, str(_FO_ROOT / "scripts"))
from check_safedir_required import (  # noqa: E402  — sys.path manipulation above
    _ALLOWLISTED_FILES,
    _FLAGGED_CALLS,
    _call_name,
    _collect_marker_comment_lines,
    _has_opt_out_in_window,
    find_violations,
    scan_tree,
)


def _synth(tmp_path: Path, content: str, name: str = "mod.py") -> Path:
    src = tmp_path / "src" / "x"
    src.mkdir(parents=True)
    target = src / name
    target.write_text(content)
    return target


# ---------------------------------------------------------------------------
# Self-test: detector flags the expected surface
# ---------------------------------------------------------------------------


class TestDetectorPositiveCases:
    """The detector must flag every call form on the watch list."""

    @pytest.mark.parametrize(
        "source",
        [
            "import fitz\nfitz.open('/foo/bar.pdf')\n",
            "from PIL import Image\nImage.open('/foo/bar.jpg')\n",
            "import shutil\nshutil.copy2('/a', '/b')\n",
            "import shutil\nshutil.move('/a', '/b')\n",
            "import shutil\nshutil.copytree('/a', '/b')\n",
            "from docx import Document\nDocument('/foo/bar.docx')\n",
            "import docx\ndocx.Document('/foo/bar.docx')\n",
            "from pptx import Presentation\nPresentation('/foo/bar.pptx')\n",
            "from pypdf import PdfReader\nPdfReader(open('/foo/bar.pdf', 'rb'))\n",
            "from openpyxl import load_workbook\nload_workbook('/foo/bar.xlsx')\n",
            "import tarfile\ntarfile.open('/foo/bar.tar')\n",
            "from py7zr import SevenZipFile\nSevenZipFile('/foo/bar.7z')\n",
            "import rarfile\nrarfile.RarFile('/foo/bar.rar')\n",
            "import zipfile\nzipfile.ZipFile('/foo/bar.zip')\n",
            "from zipfile import ZipFile\nZipFile('/foo/bar.zip')\n",
        ],
    )
    def test_flagged_call(self, tmp_path: Path, source: str) -> None:
        violations = find_violations(_synth(tmp_path, source))
        assert len(violations) == 1, f"expected 1 violation, got {violations}"


class TestDetectorOptOut:
    """Calls carrying the ``# safedir: ok`` marker are not flagged."""

    def test_trailing_comment_opt_out(self, tmp_path: Path) -> None:
        source = "import shutil\nshutil.copy2('/a', '/b')  # safedir: ok — one-shot user export\n"
        assert find_violations(_synth(tmp_path, source)) == []

    def test_preceding_comment_opt_out(self, tmp_path: Path) -> None:
        source = (
            "import shutil\n"
            "# safedir: ok — backup of app-owned state, not user-root data\n"
            "shutil.copy2('/a', '/b')\n"
        )
        assert find_violations(_synth(tmp_path, source)) == []

    def test_string_literal_marker_does_not_opt_out(self, tmp_path: Path) -> None:
        """A string containing the marker text cannot bypass the rail."""
        source = "import shutil\nmsg = \"# safedir: ok — fake marker\"\nshutil.copy2('/a', '/b')\n"
        violations = find_violations(_synth(tmp_path, source))
        assert len(violations) == 1

    def test_bare_marker_without_reason_does_not_opt_out(self, tmp_path: Path) -> None:
        """``# safedir: ok`` without a separator + reason is rejected.

        The documented contract is ``# safedir: ok — <reason>``. A bare
        marker would let drive-by exemptions slip through review.
        """
        source = "import shutil\nshutil.copy2('/a', '/b')  # safedir: ok\n"
        assert len(find_violations(_synth(tmp_path, source))) == 1

    def test_marker_with_hyphen_separator_opts_out(self, tmp_path: Path) -> None:
        """Both ``— em-dash`` and ``- hyphen`` are accepted separators."""
        source = "import shutil\nshutil.copy2('/a', '/b')  # safedir: ok - one-shot user export\n"
        assert find_violations(_synth(tmp_path, source)) == []


class TestDetectorNegativeCases:
    """The detector must NOT flag unrelated calls."""

    @pytest.mark.parametrize(
        "source",
        [
            # Read-only that doesn't follow symlinks for content
            "from pathlib import Path\nPath('/foo').exists()\n",
            # safe_walk is the safe primitive
            "from core.path_guard import safe_walk\nlist(safe_walk('/foo'))\n",
            # open() with a read mode — covered by atomic-write rail, not this one
            "open('/foo', 'r')\n",
            # logger noise
            "import logging\nlogging.info('shutil.copy2 was called')\n",
        ],
    )
    def test_not_flagged(self, tmp_path: Path, source: str) -> None:
        assert find_violations(_synth(tmp_path, source)) == []


class TestDetectorHelpers:
    """Unit-coverage on the AST helper functions."""

    def test_call_name_dotted(self, tmp_path: Path) -> None:
        import ast as _ast

        tree = _ast.parse("a.b.c(1)")
        call = next(n for n in _ast.walk(tree) if isinstance(n, _ast.Call))
        assert _call_name(call) == "a.b.c"

    def test_call_name_bare(self, tmp_path: Path) -> None:
        import ast as _ast

        tree = _ast.parse("Document('x')")
        call = next(n for n in _ast.walk(tree) if isinstance(n, _ast.Call))
        assert _call_name(call) == "Document"

    def test_call_name_dynamic_returns_none(self, tmp_path: Path) -> None:
        import ast as _ast

        tree = _ast.parse("getattr(x, 'open')(1)")
        call = next(n for n in _ast.walk(tree) if isinstance(n, _ast.Call))
        assert _call_name(call) is None

    def test_marker_window_bounds(self) -> None:
        """Marker on the call line + 6 below = inside; 7 below = outside."""
        # call at line 10; window covers 8..16 inclusive
        assert _has_opt_out_in_window({16}, call_line=10, total_lines=100) is True
        assert _has_opt_out_in_window({17}, call_line=10, total_lines=100) is False
        assert _has_opt_out_in_window({8}, call_line=10, total_lines=100) is True
        assert _has_opt_out_in_window({7}, call_line=10, total_lines=100) is False

    def test_marker_comment_only_real_comments(self) -> None:
        """A docstring or string literal containing the marker doesn't count."""
        source = (
            '"""# safedir: ok — fake"""\n'
            'x = "# safedir: ok — also fake"\n'
            "# safedir: ok — real\n"
            "y = 1\n"
        )
        marker_lines = _collect_marker_comment_lines(source)
        assert marker_lines == {3}


class TestPredicateNegativeCases:
    """T10 backfill — predicate-style detectors need negative cases that
    pass the same surface shape with wrong context and assert ``False``.

    ``_call_name`` returns the dotted attribute chain. The detector then
    matches against ``_FLAGGED_CALLS``. Two false-positive shapes worth
    verifying explicitly:
    """

    def test_attribute_chain_with_different_base_not_flagged(self, tmp_path: Path) -> None:
        """``other.copy2`` (not ``shutil.copy2``) must not match."""
        source = "class X:\n    def copy2(self, a, b): pass\nx = X()\nx.copy2('a', 'b')\n"
        # _call_name returns "x.copy2" — not in _FLAGGED_CALLS.
        assert find_violations(_synth(tmp_path, source)) == []

    def test_unrelated_open_method_not_flagged(self, tmp_path: Path) -> None:
        """A method called ``open`` on an unrelated receiver is not flagged.

        Exercises the dotted-call path: ``s.open(...)`` parses to
        ``Attribute(Name('s'), 'open')``, so ``_call_name`` returns
        ``"s.open"`` — not in ``_FLAGGED_CALLS``. Compare with calling
        ``S()`` directly, where ``_call_name`` would return ``None``
        because the base is a ``Call`` rather than a ``Name`` (that
        false-negative path is tested separately in
        ``TestDetectorHelpers.test_call_name_dynamic_returns_none``).
        """
        source = "class S:\n    def open(self, addr): pass\ns = S()\ns.open('localhost')\n"
        assert find_violations(_synth(tmp_path, source)) == []


# ---------------------------------------------------------------------------
# Live-tree advisory run
# ---------------------------------------------------------------------------


class TestLiveTreeAdvisory:
    """Run the detector against the real ``src/`` in advisory mode."""

    def test_advisory_run_exits_zero(self) -> None:
        """The rail is advisory in PR1 — must not fail CI on the current tree."""
        result = subprocess.run(
            [sys.executable, str(_SCRIPT), "--advisory"],
            capture_output=True,
            text=True,
            cwd=_FO_ROOT,
        )
        assert result.returncode == 0, (
            f"advisory rail unexpectedly failed (rc={result.returncode}):\n"
            f"stdout: {result.stdout}\nstderr: {result.stderr}"
        )

    def test_baseline_visible_in_stderr(self) -> None:
        """The advisory run prints the violation count so PR reviewers see drift.

        Phase-1 baseline: ~42 call sites across read-side, dedupe, undo, and
        watcher paths. The exact number is implementation-noise (one shuffle
        of ``shutil.copy2`` lands ±1). We assert the report is present and
        the per-callee breakdown emits something useful, not the exact count.
        """
        result = subprocess.run(
            [sys.executable, str(_SCRIPT), "--advisory"],
            capture_output=True,
            text=True,
            cwd=_FO_ROOT,
        )
        assert "[safedir-rail]" in result.stderr
        assert "call site(s)" in result.stderr
        assert "Breakdown by callee:" in result.stderr
        assert "ADVISORY mode" in result.stderr

    def test_allowlist_files_not_scanned(self) -> None:
        """Allowlisted implementation files don't generate self-flags."""
        violations = scan_tree()
        flagged_files = {p.relative_to(_FO_ROOT).as_posix() for p, _, _, _ in violations}
        for allowlisted in _ALLOWLISTED_FILES:
            assert allowlisted not in flagged_files, (
                f"allowlisted file {allowlisted} appeared in violations"
            )


class TestFlaggedCallsContract:
    """Lock in the set of detected calls so silent changes are visible."""

    def test_flagged_calls_covers_the_audited_surface(self) -> None:
        """The minimum surface flagged by phase 1 — listed in #265."""
        required = {
            "fitz.open",
            "Image.open",
            "Presentation",
            "load_workbook",
            "PdfReader",
            "tarfile.open",
            "SevenZipFile",
            "RarFile",
            "shutil.copy2",
            "shutil.move",
            "shutil.copytree",
            "zipfile.ZipFile",
            "ZipFile",
        }
        missing = required - _FLAGGED_CALLS
        assert not missing, f"watch list dropped flagged calls: {missing}"


# ---------------------------------------------------------------------------
# Bare-open detection — added in PR3g (deferred from #271)
# ---------------------------------------------------------------------------


class TestBareOpenDetectionDirScoped:
    """Bare ``open(path, "r"...)`` / ``Path.open("r"...)`` / ``io.open(path, "r"...)``
    calls are only flagged inside directories listed in
    ``_READ_OPEN_ENFORCED_DIRS``. PR3g adds the detection with an
    empty enforced-dirs set as a placeholder; PR3i populates it for
    migrated reader / dedup directories.
    """

    def test_bare_open_not_flagged_when_dir_set_empty(self, tmp_path: Path) -> None:
        """With no directories enforced, bare-open reads aren't flagged
        even in synthetic files — the detection class is opt-in.
        """
        source = "with open('/x/y.txt', 'rb') as f: pass\n"
        # Synthetic path under tmp/src/x — not in _READ_OPEN_ENFORCED_DIRS.
        # Empty set means no file is enforced.
        violations = find_violations(_synth(tmp_path, source))
        assert violations == []

    @pytest.mark.parametrize(
        "source",
        [
            "with open('/x/y.txt', 'rb') as f: pass\n",
            'with open("/x", mode="r") as f: pass\n',
            "with open('/x') as f: pass\n",  # default mode = "r"
            "import io\nwith io.open('/x', 'rb') as f: pass\n",
            "from pathlib import Path\nPath('/x').open('rb')\n",
            "from pathlib import Path\nPath('/x').open()\n",  # default mode
            "from pathlib import Path\nPath('/x').open(mode='r')\n",
        ],
    )
    def test_bare_open_detected_via_helper(self, tmp_path: Path, source: str) -> None:
        """Direct unit-test of the AST helper — bypasses the dir-scope
        gate so we can verify the detection logic itself.
        """
        from check_safedir_required import _bare_open_violation

        tree = ast.parse(source)
        calls = [n for n in ast.walk(tree) if isinstance(n, ast.Call)]
        flagged = [_bare_open_violation(c) for c in calls if _bare_open_violation(c) is not None]
        assert len(flagged) == 1, f"expected exactly one read-mode call, got {flagged}"

    @pytest.mark.parametrize(
        "source",
        [
            # Write modes are not reads.
            "with open('/x', 'wb') as f: pass\n",
            'with open("/x", mode="w") as f: pass\n',
            "with open('/x', 'a') as f: pass\n",
            "from pathlib import Path\nPath('/x').open('wb')\n",
            "from pathlib import Path\nPath('/x').open(mode='a')\n",
            "import io\nio.open('/x', 'wb')\n",
        ],
    )
    def test_write_modes_not_flagged(self, tmp_path: Path, source: str) -> None:
        """Write-only modes (``"w"``/``"a"``/``"x"`` without ``"+"``) are
        legitimate writes, not reads — handled by the atomic-write rail
        instead. The bare-open read detector must not flag them.
        """
        from check_safedir_required import _bare_open_violation

        tree = ast.parse(source)
        calls = [n for n in ast.walk(tree) if isinstance(n, ast.Call)]
        flagged = [_bare_open_violation(c) for c in calls if _bare_open_violation(c) is not None]
        assert flagged == [], f"write-mode call wrongly flagged: {flagged}"

    def test_read_plus_modes_flagged(self) -> None:
        """``"r+"`` / ``"w+"`` / ``"a+"`` reads-and-writes still open
        the underlying file and could dereference a symlink."""
        from check_safedir_required import _bare_open_violation

        for source in [
            "open('/x', 'r+')\n",
            "open('/x', 'w+')\n",
            "open('/x', 'a+')\n",
        ]:
            tree = ast.parse(source)
            call = next(n for n in ast.walk(tree) if isinstance(n, ast.Call))
            assert _bare_open_violation(call) is not None, f"missed read-write mode: {source!r}"

    @pytest.mark.parametrize(
        "source,expected_name",
        [
            # Module-style ``.open`` APIs share the ``open(file, mode)``
            # builtin signature — mode at position 1, not 0. The filename
            # at position 0 must NOT be misread as the mode.
            ("import gzip\ngzip.open('/x/a', 'rb')\n", "gzip.open"),
            ("import bz2\nbz2.open('/x/file.txt', 'rb')\n", "bz2.open"),
            ("import lzma\nlzma.open('/x/data', 'r')\n", "lzma.open"),
            ("import tarfile\ntarfile.open('/x/archive', 'r')\n", "tarfile.open"),
            ("import builtins\nbuiltins.open('/x/x', 'rb')\n", "builtins.open"),
        ],
    )
    def test_module_style_open_flagged_with_correct_mode_position(
        self, source: str, expected_name: str
    ) -> None:
        """Regression for the filename-as-mode false negative: a call like
        ``gzip.open("/x/a", "rb")`` must extract mode from arg 1, not arg
        0. Previously ``"/x/a"`` got interpreted as mode → letter ``'a'``
        intersects write-only set → not flagged (false negative)."""
        from check_safedir_required import _bare_open_violation

        tree = ast.parse(source)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call) and isinstance(n.func, ast.Attribute)
        )
        assert _bare_open_violation(call) == expected_name

    @pytest.mark.parametrize(
        "source",
        [
            # Write modes on module-style APIs must NOT be flagged.
            "import gzip\ngzip.open('/x/out', 'wb')\n",
            "import bz2\nbz2.open('/x/out', 'w')\n",
            "import lzma\nlzma.open('/x/out', 'a')\n",
            "import tarfile\ntarfile.open('/x/out', 'x')\n",
        ],
    )
    def test_module_style_open_write_modes_not_flagged(self, source: str) -> None:
        """Verifies the mode-position fix doesn't over-flag — write-only
        calls on module-style APIs still get correctly dropped."""
        from check_safedir_required import _bare_open_violation

        tree = ast.parse(source)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call) and isinstance(n.func, ast.Attribute)
        )
        assert _bare_open_violation(call) is None

    def test_filename_with_write_mode_letters_does_not_break_classification(self) -> None:
        """Belt-and-suspenders: filenames containing ``'w'`` / ``'a'`` /
        ``'x'`` letters used to mis-classify the call as write-only because
        the wrong arg was consulted. Now the mode position is correct, so
        the filename content is irrelevant."""
        from check_safedir_required import _bare_open_violation

        for filename in ["/x/awx", "writer.log", "/y/data.xml"]:
            source = f"import gzip\ngzip.open({filename!r}, 'rb')\n"
            tree = ast.parse(source)
            call = next(
                n
                for n in ast.walk(tree)
                if isinstance(n, ast.Call) and isinstance(n.func, ast.Attribute)
            )
            assert _bare_open_violation(call) == "gzip.open", (
                f"filename {filename!r} broke gzip.open detection"
            )

    @pytest.mark.parametrize(
        "source,expected_name",
        [
            # ``import gzip as gz`` — receiver name 'gz' must resolve to 'gzip'
            # via the alias map before module-style classification applies.
            ("import gzip as gz\ngz.open('/x/a', 'rb')\n", "gzip.open"),
            ("import bz2 as bz\nbz.open('/x/data', 'r')\n", "bz2.open"),
            ("import lzma as xz\nxz.open('/x/x', 'rb')\n", "lzma.open"),
            ("import tarfile as tf\ntf.open('/x/archive', 'r')\n", "tarfile.open"),
        ],
    )
    def test_aliased_module_imports_resolve_to_module_style(
        self, source: str, expected_name: str
    ) -> None:
        """Regression for Codex P2 (8f7e18d): aliased module imports fell
        through to the Path.open branch because the receiver name didn't
        match the literal module-style allowlist. The fix builds an alias
        map from ``ast.Import`` and resolves through it before checking
        the allowlist — so ``gz.open(...)`` now correctly classifies as
        ``gzip.open`` and consults mode at position 1."""
        from check_safedir_required import _bare_open_violation, _build_module_alias_map

        tree = ast.parse(source)
        aliases = _build_module_alias_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call) and isinstance(n.func, ast.Attribute)
        )
        assert _bare_open_violation(call, aliases) == expected_name

    @pytest.mark.parametrize(
        "source",
        [
            # Aliased imports with write mode must NOT be flagged — proves
            # the alias resolution doesn't introduce a false positive.
            "import gzip as gz\ngz.open('/x/out', 'wb')\n",
            "import bz2 as bz\nbz.open('/x/out', 'w')\n",
            "import lzma as xz\nxz.open('/x/out', 'a')\n",
        ],
    )
    def test_aliased_module_imports_write_modes_not_flagged(self, source: str) -> None:
        """The alias-aware module-style path must drop write-only modes the
        same way the literal-name path does."""
        from check_safedir_required import _bare_open_violation, _build_module_alias_map

        tree = ast.parse(source)
        aliases = _build_module_alias_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call) and isinstance(n.func, ast.Attribute)
        )
        assert _bare_open_violation(call, aliases) is None

    def test_build_module_alias_map_handles_plain_and_aliased(self) -> None:
        """Unit-test the alias map builder directly: plain ``import X`` is
        the identity mapping, ``import X as Y`` records ``Y → X``, and
        dotted imports collapse to the top-level module name (matching
        the AST receiver shape for ``X.open(...)``)."""
        from check_safedir_required import _build_module_alias_map

        source = (
            "import gzip\n"
            "import bz2 as bz\n"
            "import tarfile as tf\n"
            "import lzma\n"
            "import pathlib.PurePath  # not a real stmt but parses\n"
        )
        # Note: ``import a.b`` is real syntax; the local name bound is the
        # top-level (``a``). Verify that.
        source = (
            "import gzip\n"
            "import bz2 as bz\n"
            "import tarfile as tf\n"
            "import lzma\n"
            "import collections.abc\n"
        )
        tree = ast.parse(source)
        aliases = _build_module_alias_map(tree)
        assert aliases["gzip"] == "gzip"
        assert aliases["bz"] == "bz2"
        assert aliases["tf"] == "tarfile"
        assert aliases["lzma"] == "lzma"
        # `import collections.abc` binds `collections` locally; the alias
        # map collapses to the top-level module name.
        assert aliases["collections"] == "collections"

    @pytest.mark.parametrize(
        "source",
        [
            # Module-level rebind BEFORE the call → alias shadowed.
            "import gzip as gz\ngz = object()\ngz.open('wb')\n",
            # Path() rebinding (Codex's first example) before the call.
            ("from pathlib import Path\nimport gzip as gz\ngz = Path('/x')\ngz.open('wb')\n"),
            # For-loop rebind before the call.
            "import gzip as gz\nfor gz in []: pass\ngz.open('wb')\n",
            # With-as rebind before the call.
            "import gzip as gz\nwith open('/x') as gz: pass\ngz.open('wb')\n",
            # Walrus rebind before the call.
            "import gzip as gz\nif (gz := object()): pass\ngz.open('wb')\n",
        ],
    )
    def test_alias_shadowed_when_module_rebind_precedes_call(self, source: str) -> None:
        """Regression for Codex P2 (cf2c841 + earlier): a module-level
        rebind that happens BEFORE the call shadows the alias at the
        call site. Verified through ``_active_aliases_at_call`` which
        is what ``find_violations`` actually uses. ``_build_module_alias_map``
        no longer filters — it returns all imports — so the per-call
        resolver decides effective aliases with scope + order awareness."""
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        # Build phase keeps the alias — no filter at this layer anymore.
        assert base_aliases.get("gz") == "gzip"
        parents = _build_parent_map(tree)
        # The Attribute-style ``gz.open(...)`` call is what we care about.
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        assert "gz" not in active

    def test_unshadowed_alias_kept_in_resolver(self) -> None:
        """Belt-and-suspenders for the resolver: a non-shadowed alias
        survives the per-call resolution."""
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        source = "import gzip as gz\ngz.open('/x/a', 'rb')\n"
        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        assert active.get("gz") == "gzip"

    def test_alias_active_when_rebind_is_after_call(self) -> None:
        """Codex's fresh P2 example (8a73104): a call that comes BEFORE
        a later module-level rebind must keep the alias active. Without
        order awareness, the file-global filter wrongly dropped the
        alias and the call fell through to Path.open with the filename
        misread as mode — masking real reads in enforced directories.

        Source under test::

            import gzip as gz
            gz.open('/x/ax', 'rb')   # ← real READ; alias still active
            gz = object()            # ← rebind AFTER the call
        """
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        source = "import gzip as gz\ngz.open('/x/ax', 'rb')\ngz = object()\n"
        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        # Rebind is AFTER the call → alias still active → resolver
        # returns the canonical mapping.
        assert active.get("gz") == "gzip"

    @pytest.mark.parametrize(
        "source",
        [
            # Function parameter shadows the alias INSIDE the function.
            "import gzip as gz\ndef f(gz):\n    gz.open('wb')\n",
            # Lambda parameter — inner scope.
            "import gzip as gz\nh = lambda gz: gz.open('wb')\n",
            # Keyword-only parameter.
            "import gzip as gz\ndef f(*, gz):\n    gz.open('wb')\n",
            # Positional-only parameter.
            "import gzip as gz\ndef f(gz, /):\n    gz.open('wb')\n",
            # Async function parameter.
            "import gzip as gz\nasync def f(gz):\n    gz.open('wb')\n",
            # Function-body assignment (not parameter) — Python's
            # function-scope rule binds it locally regardless of order.
            "import gzip as gz\ndef f():\n    gz = object()\n    gz.open('wb')\n",
        ],
    )
    def test_alias_dropped_when_call_is_in_function_with_local_binding(self, source: str) -> None:
        """Function-scope shadowing for calls INSIDE the function.

        ``def f(gz): gz.open('wb')`` — the call inside ``f`` refers to
        the parameter, not the module alias. Resolver must drop the
        alias at this call site even though the module-level alias is
        otherwise active."""
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        assert "gz" not in active

    def test_alias_kept_for_module_call_when_function_locally_shadows(self) -> None:
        """Symmetric to the function-shadow test: a function whose body
        locally shadows ``gz`` must NOT shadow ``gz`` for calls OUTSIDE
        that function. Without proper scope awareness, my prior file-
        global filter dropped the alias for both calls."""
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        source = (
            "import gzip as gz\n"
            "def f(gz):\n"
            "    gz.open('wb')\n"  # function-local — alias shadowed
            "gz.open('/x/a', 'rb')\n"  # module-level — alias active
        )
        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        # Collect both Attribute-style ``.open`` calls.
        open_calls = [
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        ]
        # Inside the function (lineno 3) → alias shadowed.
        inside_call = next(c for c in open_calls if c.lineno == 3)
        # At module level (lineno 4) → alias active.
        module_call = next(c for c in open_calls if c.lineno == 4)
        assert "gz" not in _active_aliases_at_call(inside_call, base_aliases, parents, tree)
        assert _active_aliases_at_call(module_call, base_aliases, parents, tree).get("gz") == "gzip"

    def test_nested_scope_import_does_not_clobber_module_alias(self) -> None:
        """Regression for Codex P1 (2fdb63e): an import inside a function
        body binds its name in the function's local scope, NOT at module
        level. If ``_build_module_alias_map`` walked the full tree, a
        nested ``import pathlib as gz`` would clobber the module-level
        ``import gzip as gz`` mapping and hide the real read at line 2.

        Source under test::

            import gzip as gz
            gz.open('/x/a', 'rb')            # ← module-level read
            def helper(): import pathlib as gz  # ← nested import, NOT a module alias
        """
        from check_safedir_required import _build_module_alias_map

        source = (
            "import gzip as gz\ngz.open('/x/a', 'rb')\ndef helper():\n    import pathlib as gz\n"
        )
        tree = ast.parse(source)
        aliases = _build_module_alias_map(tree)
        # Module-level alias unaffected by nested-scope import.
        assert aliases.get("gz") == "gzip"

    @pytest.mark.parametrize(
        "source",
        [
            # ``import X as gz`` inside the function body.
            "import gzip as gz\ndef f():\n    import pathlib as gz\n    gz.open('wb')\n",
            # ``import gz`` inside the function (no asname; binds top-level
            # name ``gz`` — same shadow effect).
            "import gzip as gz\ndef f():\n    import gz  # noqa\n    gz.open('wb')\n",
            # ``from pathlib import Path as gz`` inside function.
            "import gzip as gz\ndef f():\n    from pathlib import Path as gz\n    gz.open('wb')\n",
        ],
    )
    def test_function_scope_import_shadows_module_alias(self, source: str) -> None:
        """Regression for Codex P2 (7c59f11): an ``import ... as gz``
        inside a function body shadows the module-level alias for calls
        within that function. ``_function_local_names`` must collect
        ``ast.Import`` / ``ast.ImportFrom`` alias names, not just
        ``Name(Store)``."""
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        # In-function import shadows module-level alias.
        assert "gz" not in active

    @pytest.mark.parametrize(
        "source",
        [
            # for-target shadow — single-line.
            "import gzip as gz\nfor gz in []: gz.open('wb')\n",
            # for-target shadow — multi-line.
            "import gzip as gz\nfor gz in []:\n    gz.open('wb')\n",
            # with-target shadow.
            "import gzip as gz\nwith open('/x') as gz: gz.open('wb')\n",
        ],
    )
    def test_alias_dropped_by_same_line_header_binding(self, source: str) -> None:
        """Regression for Codex P2 (1f6e27a): the previous order check
        used ``node.lineno < cutoff`` (strict less-than), which missed
        bindings on the SAME line as the cutoff statement — exactly the
        case for ``for gz in ...:`` and ``with ... as gz:`` where the
        target binding and the call share the statement's lineno. The
        replay model walks all statements with ``lineno <= cutoff`` and
        collects in-header bindings before evaluating the call."""
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        assert "gz" not in active

    def test_replay_preserves_earlier_import_for_calls_before_second_import(
        self,
    ) -> None:
        """Regression for Codex P2 (1f6e27a): when two module-level
        imports bind the same local name (``import gzip as gz`` then
        ``import pathlib as gz``), a call BETWEEN them sees the
        earlier mapping; a call AFTER the second sees the new one.

        Source under test::

            import gzip as gz         # 1
            gz.open('/x/a', 'rb')     # 2  ← resolves to gzip.open
            import pathlib as gz      # 3
            gz.open('rb')             # 4  ← resolves to pathlib.open

        Both calls have ``gz`` as the receiver, but the replay model
        gives them different alias maps.
        """
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        source = "import gzip as gz\ngz.open('/x/a', 'rb')\nimport pathlib as gz\ngz.open('rb')\n"
        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        calls = [
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        ]
        # Call at line 2 — before the second import.
        first_call = next(c for c in calls if c.lineno == 2)
        # Call at line 4 — after the second import.
        second_call = next(c for c in calls if c.lineno == 4)
        first_active = _active_aliases_at_call(first_call, base_aliases, parents, tree)
        second_active = _active_aliases_at_call(second_call, base_aliases, parents, tree)
        assert first_active.get("gz") == "gzip"
        assert second_active.get("gz") == "pathlib"

    @pytest.mark.parametrize(
        "source",
        [
            # List comprehension target.
            "import gzip as gz\n[gz for gz in []]\ngz.open('/x/a', 'rb')\n",
            # Set comprehension.
            "import gzip as gz\n{gz for gz in []}\ngz.open('/x/a', 'rb')\n",
            # Dict comprehension (key + value).
            "import gzip as gz\n{gz: gz for gz in []}\ngz.open('/x/a', 'rb')\n",
            # Generator expression.
            "import gzip as gz\nlist(gz for gz in [])\ngz.open('/x/a', 'rb')\n",
        ],
    )
    def test_comprehension_targets_do_not_shadow_module_alias(self, source: str) -> None:
        """Regression for Codex P1 (d6e2911): Python 3 scopes
        comprehension/genexpr targets to the comprehension itself —
        they do NOT shadow names in the enclosing scope. ``[gz for gz
        in xs]`` followed by ``gz.open('/x/a', 'rb')`` must keep the
        module alias active for the later call."""
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        # The Attribute-style call (gz.open) — not the comprehension internals.
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        # Alias survives the comprehension.
        assert active.get("gz") == "gzip"

    def test_class_body_local_shadows_alias_for_direct_class_call(self) -> None:
        """Regression for Codex P2 (d6e2911): a name bound in a class
        body (e.g. ``gz = Writer()``) shadows the module-level alias
        for calls placed DIRECTLY in the class body (not inside a
        method).

        Source under test::

            import gzip as gz
            class C:
                gz = Writer()
                gz.open('wb')   # ← class-body call — gz is class-local
        """
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        source = "import gzip as gz\nclass C:\n    gz = object()\n    gz.open('wb')\n"
        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        # Class body shadows.
        assert "gz" not in active

    def test_method_call_does_not_see_class_body_shadow(self) -> None:
        """Symmetric to the class-body test: a call inside a METHOD does
        NOT see class-body bindings (Python's LEGB rule skips the class
        namespace for method-internal name resolution).

        Source under test::

            import gzip as gz
            class C:
                gz = Writer()
                def m(self):
                    gz.open('/x/a', 'rb')   # ← method call uses MODULE-level gz
        """
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        source = (
            "import gzip as gz\n"
            "class C:\n"
            "    gz = object()\n"
            "    def m(self):\n"
            "        gz.open('/x/a', 'rb')\n"
        )
        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        # Class body's ``gz`` does NOT leak into method scope. Method
        # name resolution skips class to module. Alias stays active.
        assert active.get("gz") == "gzip"

    @pytest.mark.parametrize(
        "source",
        [
            # ``from M import X as gz`` rebinds gz to a class/function.
            "import gzip as gz\nfrom pathlib import Path as gz\ngz.open('wb')\n",
            # ``from M import gz`` (no asname) also rebinds.
            "import gzip as gz\nfrom mod import gz\ngz.open('wb')\n",
            # ``def gz(...)`` rebinds.
            "import gzip as gz\ndef gz(): pass\ngz.open('wb')\n",
            # ``async def gz(...)`` rebinds.
            "import gzip as gz\nasync def gz(): pass\ngz.open('wb')\n",
            # ``class gz: ...`` rebinds.
            "import gzip as gz\nclass gz: pass\ngz.open('wb')\n",
        ],
    )
    def test_module_scope_non_assign_rebinds_drop_alias(self, source: str) -> None:
        """Regression for Codex P2 (308cd7b): module-scope rebinders
        other than Name(Store) — ``from ... import ... as``, ``def``,
        ``async def``, ``class`` — must also drop the alias. Before
        this fix, ``import gzip as gz; from pathlib import Path as gz;
        gz.open('wb')`` kept the gz → gzip mapping, mis-classified the
        call as module-style ``gzip.open`` (mode at arg 1 → omitted →
        default-read), and would have flagged a legitimate write-only
        ``Path.open``."""
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(call, base_aliases, parents, tree)
        assert "gz" not in active

    def test_in_function_call_drops_alias_on_any_module_rebind(self) -> None:
        """Regression for Codex P2 (2fdb63e): a call inside a function
        sees the *runtime* value of the module global by the time the
        function is invoked. So ANY module-level rebind anywhere in the
        file shadows the alias for in-function calls — not just rebinds
        before the function definition.

        Source under test::

            import gzip as gz
            def f():
                gz.open('wb')         # ← in-function call
            gz = Path('/x/out')      # ← rebind AFTER def
        """
        from check_safedir_required import (
            _active_aliases_at_call,
            _build_module_alias_map,
            _build_parent_map,
        )

        source = (
            "from pathlib import Path\n"
            "import gzip as gz\n"
            "def f():\n"
            "    gz.open('wb')\n"
            "gz = Path('/x/out')\n"
        )
        tree = ast.parse(source)
        base_aliases = _build_module_alias_map(tree)
        parents = _build_parent_map(tree)
        in_func_call = next(
            n
            for n in ast.walk(tree)
            if isinstance(n, ast.Call)
            and isinstance(n.func, ast.Attribute)
            and n.func.attr == "open"
        )
        active = _active_aliases_at_call(in_func_call, base_aliases, parents, tree)
        # Rebind happens after the def, but the function call sees the
        # post-rebind value at invocation time → alias dropped.
        assert "gz" not in active

    def test_find_violations_respects_parameter_shadowing_end_to_end(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Full-stack regression for Codex P2 (parameter shadowing): a
        file under an enforced dir that uses ``gz`` as a function
        parameter must NOT classify ``gz.open('wb')`` as module-style.
        Codex's exact example."""
        import check_safedir_required as _csr

        src_root = tmp_path / "src"
        src_root.mkdir()
        target_dir = src_root / "myreaders"
        target_dir.mkdir()
        target = target_dir / "mod.py"
        target.write_text("import gzip as gz\ndef f(gz):\n    gz.open('wb')\n")

        monkeypatch.setattr(_csr, "_SRC_DIR", src_root)
        monkeypatch.setattr(_csr, "_READ_OPEN_ENFORCED_DIRS", frozenset({"src/myreaders"}))
        violations = _csr.find_violations(target)

        # ``gz`` is a function parameter → dropped from alias map →
        # falls to Path.open branch → mode at arg 0 = 'wb' → write-only
        # → not flagged.
        assert violations == []


class TestBareOpenDetectionDirScopedIntegration:
    """Integration test: when ``_READ_OPEN_ENFORCED_DIRS`` is non-empty,
    ``find_violations`` reports bare reads on files inside the enforced
    directory. Verifies the pass-2 wiring + ``_file_under_enforced_dir``
    end-to-end, not just the helper in isolation.

    Patches the module-level ``_READ_OPEN_ENFORCED_DIRS`` set so the
    test owns the membership for its duration; the patch is reverted
    after the test runs.
    """

    def test_find_violations_flags_bare_open_when_dir_enforced(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import check_safedir_required as _csr

        # Create the same ``src/x/<file>`` layout ``_synth`` uses so the
        # ``_file_under_enforced_dir`` resolution succeeds: it computes
        # the path relative to ``_SRC_DIR``'s parent. We must point
        # ``_SRC_DIR`` at our synthetic root so its parent is tmp_path.
        src_root = tmp_path / "src"
        src_root.mkdir()
        target_dir = src_root / "myreaders"
        target_dir.mkdir()
        target = target_dir / "mod.py"
        target.write_text("with open('/x/y.txt', 'rb') as f: pass\n")

        monkeypatch.setattr(_csr, "_SRC_DIR", src_root)
        monkeypatch.setattr(_csr, "_READ_OPEN_ENFORCED_DIRS", frozenset({"src/myreaders"}))
        violations = _csr.find_violations(target)

        # The bare ``open(path, "rb")`` is now flagged because the
        # directory is enforced.
        assert len(violations) == 1
        line_no, name, _excerpt = violations[0]
        assert name == "open"
        assert line_no == 1

    def test_find_violations_skips_bare_open_outside_enforced_dir(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """File NOT under an enforced dir → bare open ignored even when
        the set is populated."""
        import check_safedir_required as _csr

        src_root = tmp_path / "src"
        src_root.mkdir()
        target_dir = src_root / "elsewhere"
        target_dir.mkdir()
        target = target_dir / "mod.py"
        target.write_text("with open('/x/y.txt', 'rb') as f: pass\n")

        monkeypatch.setattr(_csr, "_SRC_DIR", src_root)
        # Different directory enforced.
        monkeypatch.setattr(_csr, "_READ_OPEN_ENFORCED_DIRS", frozenset({"src/myreaders"}))
        violations = _csr.find_violations(target)

        assert violations == []

    def test_find_violations_respects_opt_out_marker_for_bare_open(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Trailing ``# safedir: ok — <reason>`` marker exempts a bare
        open just like it does for library-call violations."""
        import check_safedir_required as _csr

        src_root = tmp_path / "src"
        src_root.mkdir()
        target_dir = src_root / "myreaders"
        target_dir.mkdir()
        target = target_dir / "mod.py"
        target.write_text(
            "with open('/x/y.txt', 'rb') as f: pass  # safedir: ok — legacy callsite\n"
        )

        monkeypatch.setattr(_csr, "_SRC_DIR", src_root)
        monkeypatch.setattr(_csr, "_READ_OPEN_ENFORCED_DIRS", frozenset({"src/myreaders"}))
        violations = _csr.find_violations(target)

        assert violations == []

    def test_find_violations_resolves_aliased_module_import_end_to_end(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Full-stack regression for Codex P2: a file under an enforced dir
        with ``import gzip as gz`` followed by ``gz.open(path, 'rb')`` must
        be flagged. Before the alias-map fix, the bare-open detector
        consulted only the literal receiver id (``gz``), missed the
        allowlist, and fell through to the Path.open branch — which then
        misread the filename as the mode."""
        import check_safedir_required as _csr

        src_root = tmp_path / "src"
        src_root.mkdir()
        target_dir = src_root / "myreaders"
        target_dir.mkdir()
        target = target_dir / "mod.py"
        target.write_text("import gzip as gz\ngz.open('/x/a', 'rb')\n")

        monkeypatch.setattr(_csr, "_SRC_DIR", src_root)
        monkeypatch.setattr(_csr, "_READ_OPEN_ENFORCED_DIRS", frozenset({"src/myreaders"}))
        violations = _csr.find_violations(target)

        assert len(violations) == 1
        line_no, name, _excerpt = violations[0]
        assert name == "gzip.open"  # canonical name, not literal alias
        assert line_no == 2

    def test_find_violations_aliased_write_mode_not_flagged_end_to_end(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Companion to the previous test — same alias setup but write
        mode. Confirms the alias fix doesn't over-flag write-only calls."""
        import check_safedir_required as _csr

        src_root = tmp_path / "src"
        src_root.mkdir()
        target_dir = src_root / "myreaders"
        target_dir.mkdir()
        target = target_dir / "mod.py"
        target.write_text("import gzip as gz\ngz.open('/x/out', 'wb')\n")

        monkeypatch.setattr(_csr, "_SRC_DIR", src_root)
        monkeypatch.setattr(_csr, "_READ_OPEN_ENFORCED_DIRS", frozenset({"src/myreaders"}))
        violations = _csr.find_violations(target)

        assert violations == []

    def test_find_violations_respects_alias_shadowing_end_to_end(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Full-stack regression for Codex P2 (scope awareness): a file
        that imports ``gzip as gz`` but then rebinds ``gz`` to a Path
        before calling ``.open('wb')`` must NOT be flagged. The
        rebinding signals that the call is Path-style — mode at arg 0
        is ``'wb'``, write-only, not a read. Without scope awareness,
        the file-global alias map would module-style-classify the call
        and incorrectly flag it as a read (mode at arg 1 → omitted →
        default read)."""
        import check_safedir_required as _csr

        src_root = tmp_path / "src"
        src_root.mkdir()
        target_dir = src_root / "myreaders"
        target_dir.mkdir()
        target = target_dir / "mod.py"
        target.write_text(
            "from pathlib import Path\nimport gzip as gz\ngz = Path('/x/out')\ngz.open('wb')\n"
        )

        monkeypatch.setattr(_csr, "_SRC_DIR", src_root)
        monkeypatch.setattr(_csr, "_READ_OPEN_ENFORCED_DIRS", frozenset({"src/myreaders"}))
        violations = _csr.find_violations(target)

        assert violations == []
