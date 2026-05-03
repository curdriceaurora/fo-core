"""Integration tests lifting six tail modules from 70-74% to >=75% line+branch coverage.

Each test class targets one module. Tests aim at specific missing line ranges
identified by ``bash .claude/scripts/measure-integration-coverage.sh``.
Closes the "Integration coverage floors" row of docs/release/beta-criteria.md
section 2 by raising every per-module floor to >=75%.
"""

from __future__ import annotations

import json
import threading
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest

pytestmark = pytest.mark.integration


# ---------------------------------------------------------------------------
# numbering.py: 74% -> 75% (+1pp)
# Targets missing lines 113-116, 186-188, 216-223, 253-261, 462-465, 488-497.
# ---------------------------------------------------------------------------


class TestNumberingLift:
    """Lift integration coverage on src/methodologies/johnny_decimal/numbering.py."""

    def _make(self) -> Any:
        from methodologies.johnny_decimal.categories import (
            AreaDefinition,
            CategoryDefinition,
            NumberingScheme,
        )
        from methodologies.johnny_decimal.numbering import JohnnyDecimalGenerator

        scheme = NumberingScheme(name="Lift", description="Lift")
        scheme.add_area(AreaDefinition(10, 19, "Finance", "Finance"))
        scheme.add_area(AreaDefinition(20, 29, "Admin", "Admin"))
        scheme.add_category(
            CategoryDefinition(area=10, category=1, name="Budgets", description="Budgets")
        )
        return scheme, JohnnyDecimalGenerator(scheme)

    def test_get_next_available_area_returns_preferred_when_free(self) -> None:
        """Hits lines 113-116 (preferred-area happy path)."""
        _, gen = self._make()
        result = gen.get_next_available_area(preferred_area=10)
        assert result == 10

    def test_generate_area_number_uses_preferred_when_free(self) -> None:
        """Hits lines 186-188 (preferred_area returns immediately)."""
        _, gen = self._make()
        num = gen.generate_area_number(name="Finance", preferred_area=10)
        assert num.area == 10
        assert num.name == "Finance"
        assert num.category is None

    def test_generate_category_number_uses_preferred_when_free(self) -> None:
        """Hits lines 216-223 (preferred_category returns immediately)."""
        _, gen = self._make()
        num = gen.generate_category_number(area=10, name="Receipts", preferred_category=5)
        assert num.area == 10
        assert num.category == 5

    def test_generate_id_number_uses_preferred_when_free(self) -> None:
        """Hits lines 253-261 (preferred_id returns immediately)."""
        _, gen = self._make()
        num = gen.generate_id_number(area=10, category=1, name="Q1", preferred_id=42)
        assert num.area == 10
        assert num.category == 1
        assert num.item_id == 42

    def test_resolve_conflict_increment_strategy_at_area_level(self, tmp_path: Path) -> None:
        """Increment strategy at the area level dispatches to ``generate_area_number`` (lines 462-465).

        Area-level resolve honors the preferred-area-free invariant: the bare
        area number is occupied but categories under it are still free, so the
        resolver returns a JohnnyDecimalNumber with the same area and the
        original name.
        """
        from methodologies.johnny_decimal.categories import JohnnyDecimalNumber

        _, gen = self._make()
        conflict = JohnnyDecimalNumber(area=10, name="X", description="d")
        gen.register_existing_number(conflict, tmp_path / "x_increment")
        resolved = gen.resolve_conflict(conflict, strategy="increment")
        assert resolved is not None
        assert resolved.name == "X"
        # Area 10 has free categories, so the resolver must reuse area 10
        # rather than wandering to area 20. category None confirms we landed in
        # the area-level branch (lines 462-465), not category- or id-level.
        assert resolved.area == 10
        assert resolved.category is None

    def test_resolve_conflict_skip_strategy_at_area_level(self, tmp_path: Path) -> None:
        """Skip strategy at the area level falls back through ``get_next_available_area`` (lines 488-497)."""
        from methodologies.johnny_decimal.categories import JohnnyDecimalNumber

        _, gen = self._make()
        conflict = JohnnyDecimalNumber(area=10, name="X")
        gen.register_existing_number(conflict, tmp_path / "x_skip")
        resolved = gen.resolve_conflict(conflict, strategy="skip")
        assert resolved is not None
        assert resolved.name == "X"
        # Same invariant as the increment case: area 10 still has free
        # categories, so skip-at-area-level returns area 10.
        assert resolved.area == 10
        assert resolved.category is None


# ---------------------------------------------------------------------------
# epub_enhanced.py: 74% -> 75% (+1pp)
# Targets missing lines 285, 287, 394-400, 438-449.
# ---------------------------------------------------------------------------


class TestEpubEnhancedLift:
    """Lift integration coverage on src/utils/epub_enhanced.py."""

    def test_clean_isbn_strips_whitespace_and_dashes(self) -> None:
        """Hits the _clean_isbn helper used at line 285."""
        pytest.importorskip("ebooklib")
        pytest.importorskip("bs4")
        from utils.epub_enhanced import EnhancedEPUBReader

        reader = EnhancedEPUBReader()
        assert reader._clean_isbn("978-3-16-148410-0") == "9783161484100"
        assert reader._clean_isbn(" isbn: 9780306406157 ") == "9780306406157"

    def test_extract_metadata_handles_uuid_identifier(self) -> None:
        """Hits line 287 (UUID branch in _extract_identifiers)."""
        pytest.importorskip("ebooklib")
        pytest.importorskip("bs4")
        from utils.epub_enhanced import EnhancedEPUBReader

        reader = EnhancedEPUBReader()
        # Build a minimal book-like object whose get_metadata returns a
        # UUID-tagged identifier.
        book = MagicMock()

        def get_metadata(ns: str, key: str) -> list[Any]:
            if ns == "DC" and key == "identifier":
                return [("urn:uuid:abc-123", {"scheme": "uuid"})]
            if ns == "DC" and key == "title":
                return [("T", {})]
            return []

        book.get_metadata = MagicMock(side_effect=get_metadata)
        book.get_items_of_type = MagicMock(return_value=[])
        book.toc = []
        book.spine = []
        book.metadata = {}

        # _extract_metadata expects the book interface; calling it directly
        # exercises the identifier-extraction branches.
        meta = reader._extract_metadata(book)
        assert meta.identifiers.get("uuid") == "urn:uuid:abc-123"

    def test_extract_chapter_title_falls_back_to_filename(self) -> None:
        """Hits lines 394-400 (filename fallback in _extract_chapter_title)."""
        pytest.importorskip("ebooklib")
        pytest.importorskip("bs4")
        from bs4 import BeautifulSoup

        from utils.epub_enhanced import EnhancedEPUBReader

        reader = EnhancedEPUBReader()
        # Soup with no headings forces fallback to filename.
        soup = BeautifulSoup("<html><body><p>No headings here</p></body></html>", "html.parser")
        item = MagicMock()
        item.file_name = "the_first-chapter.xhtml"
        title = reader._extract_chapter_title(soup, item)
        assert title == "The First Chapter"


# ---------------------------------------------------------------------------
# parallel/processor.py: 71% -> 75% (+4pp)
# Targets missing lines 117, 133-135, 237-238, 397-432.
# ---------------------------------------------------------------------------


class TestParallelProcessorLift:
    """Lift integration coverage on src/parallel/processor.py."""

    def test_process_batch_with_empty_files_short_circuits(self, tmp_path: Path) -> None:
        """Hits line 117 (early-break when remaining is empty)."""
        from parallel.config import ParallelConfig
        from parallel.processor import ParallelProcessor

        proc = ParallelProcessor(ParallelConfig(max_workers=1))
        result = proc.process_batch([], lambda p: "ok")
        assert result.succeeded == 0
        assert result.failed == 0

    def test_process_batch_consumes_full_retry_budget_on_retryable_failures(
        self, tmp_path: Path
    ) -> None:
        """Hits lines 115-120 (retry-loop iteration without non-retryable break). Non-retryable is reserved for executor-level events like timeouts; plain raised exceptions ARE retryable, so the loop iterates the full 1 + retry_count attempts."""
        from parallel.config import ParallelConfig
        from parallel.processor import ParallelProcessor

        f = tmp_path / "x.txt"
        f.write_text("data")

        attempts = {"n": 0}

        def fn(path: Path) -> str:
            attempts["n"] += 1
            raise ValueError("retryable")

        proc = ParallelProcessor(ParallelConfig(max_workers=1, retry_count=2))
        result = proc.process_batch([f], fn)
        assert result.failed == 1
        # 1 initial + 2 retries = 3 attempts.
        assert attempts["n"] == 3

    def test_process_batch_iter_with_timeout_triggers_abort(self, tmp_path: Path) -> None:
        """Slow worker exceeds ``timeout_per_file`` and trips abort/timeout branches (lines 336-355, 397-432).

        Uses a never-set ``threading.Event`` instead of ``time.sleep`` so the
        worker blocks deterministically until the test's teardown signals it,
        and the wait returns promptly when set — avoiding wall-clock flakiness
        on loaded CI runners.
        """
        from parallel.config import ParallelConfig
        from parallel.processor import ParallelProcessor

        files = [tmp_path / f"file{i}.txt" for i in range(4)]
        for f in files:
            f.write_text("data")

        blocker = threading.Event()

        def slow(path: Path) -> str:
            # blocker is never set during the test body; the wait deadline
            # (0.5s, well above timeout_per_file=0.05s) ensures workers block
            # past the timeout without burning a real wall-clock sleep.
            blocker.wait(timeout=0.5)
            return "ok"

        cfg = ParallelConfig(max_workers=1, timeout_per_file=0.05, retry_count=0)
        proc = ParallelProcessor(cfg)

        try:
            results = list(proc.process_batch_iter(files, slow))
        finally:
            # Release any still-blocked workers so the executor can shut down.
            blocker.set()

        # All files should have an error result (timed out OR aborted).
        assert len(results) >= 1
        assert all(r.success is False for r in results)
        timeout_or_abort = [
            r
            for r in results
            if "timed out" in (r.error or "").lower() or "aborted" in (r.error or "").lower()
        ]
        assert len(timeout_or_abort) >= 1

    def test_collect_result_wraps_exception(self, tmp_path: Path) -> None:
        """Hits lines 237-238 (_collect_result exception handling path)."""
        from concurrent.futures import ThreadPoolExecutor

        from parallel.processor import ParallelProcessor

        def boom() -> Any:
            raise RuntimeError("kaboom")

        with ThreadPoolExecutor(max_workers=1) as ex:
            future = ex.submit(boom)
            res = ParallelProcessor._collect_result(future, tmp_path / "p.txt")

        assert res.success is False
        assert res.error is not None
        assert "kaboom" in str(res.error)


# ---------------------------------------------------------------------------
# deduplication/extractor.py: 71% -> 75% (+4pp)
# Targets missing lines 55, 57, 168-182.
# ---------------------------------------------------------------------------


class TestDedupExtractorLift:
    """Lift integration coverage on src/services/deduplication/extractor.py."""

    def test_extract_pdf_dispatch(self, tmp_path: Path) -> None:
        """PDF extension routes through ``_extract_pdf`` (line 55 dispatch branch).

        Uses a minimal PDF that pypdf accepts but that yields no extractable
        text, so the empty-string return confirms the dispatcher chose the PDF
        branch and ran it to completion without raising.
        """
        pytest.importorskip("pypdf")
        from services.deduplication.extractor import DocumentExtractor

        # Minimal valid PDF byte sequence.
        pdf = b"%PDF-1.4\n1 0 obj <<>> endobj\ntrailer <<>>\n%%EOF\n"
        f = tmp_path / "tiny.pdf"
        f.write_bytes(pdf)
        text = DocumentExtractor().extract_text(f)
        # The PDF has no text streams, so the dispatcher returns "" — exactly
        # equal, not just "is a string".
        assert text == ""

    def test_extract_docx_with_paragraphs_and_tables(self, tmp_path: Path) -> None:
        """Hits lines 168-182 (DOCX paragraph + table extraction)."""
        pytest.importorskip("docx")
        import docx

        from services.deduplication.extractor import DocumentExtractor

        path = tmp_path / "doc.docx"
        d = docx.Document()
        d.add_paragraph("Header paragraph")
        table = d.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "CellA"
        table.cell(0, 1).text = "CellB"
        d.save(str(path))

        text = DocumentExtractor().extract_text(path)
        assert "Header paragraph" in text
        assert "CellA" in text
        assert "CellB" in text

    def test_extract_unsupported_format_raises(self, tmp_path: Path) -> None:
        """Hits line 51 (unsupported-format raise)."""
        from services.deduplication.extractor import DocumentExtractor

        f = tmp_path / "data.xyz"
        f.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported format"):
            DocumentExtractor().extract_text(f)

    def test_extract_text_missing_file_raises(self, tmp_path: Path) -> None:
        """Hits line 46 (file-not-found raise)."""
        from services.deduplication.extractor import DocumentExtractor

        with pytest.raises(OSError, match="File not found"):
            DocumentExtractor().extract_text(tmp_path / "ghost.txt")


# ---------------------------------------------------------------------------
# intelligence/profile_exporter.py: 70% -> 75% (+5pp)
# Targets missing lines 86-88, 95-97, 158-161, 216-217, 226-233, 248-250, 288-290.
# ---------------------------------------------------------------------------


class TestProfileExporterLift:
    """Lift integration coverage on src/services/intelligence/profile_exporter.py."""

    def _exporter(self) -> tuple[Any, Any]:
        from services.intelligence.profile_exporter import ProfileExporter
        from services.intelligence.profile_manager import ProfileManager

        pm = MagicMock(spec=ProfileManager)
        return ProfileExporter(pm), pm

    def _profile(self, **overrides: Any) -> Any:
        from services.intelligence.profile_manager import Profile

        prof = MagicMock(spec=Profile)
        prof.profile_name = overrides.get("name", "Test")
        prof.profile_version = overrides.get("version", "1.0")
        prof.description = overrides.get("description", "desc")
        prof.created = overrides.get("created", "2026-01-01")
        prof.updated = overrides.get("updated", "2026-05-01")
        prof.preferences = overrides.get(
            "preferences",
            {"global": {"a": 1}, "directory_specific": {"d": {"a": 2}}},
        )
        prof.learned_patterns = overrides.get("learned_patterns", {"p": "v"})
        prof.confidence_data = overrides.get("confidence_data", {"c": 0.9})
        prof.validate = MagicMock(return_value=overrides.get("valid", True))
        prof.to_dict = MagicMock(
            return_value={
                "profile_name": prof.profile_name,
                "profile_version": prof.profile_version,
                "preferences": prof.preferences,
            }
        )
        return prof

    def test_export_profile_returns_false_when_profile_missing(self, tmp_path: Path) -> None:
        """Hits lines 60-62 (profile-not-found path)."""
        exporter, pm = self._exporter()
        pm.get_profile = MagicMock(return_value=None)
        result = exporter.export_profile("ghost", tmp_path / "out.json")
        assert result is False

    def test_export_profile_returns_false_when_validation_fails(self, tmp_path: Path) -> None:
        """Hits lines 65-67 (validation-fails path)."""
        exporter, pm = self._exporter()
        pm.get_profile = MagicMock(return_value=self._profile(valid=False))
        result = exporter.export_profile("Test", tmp_path / "out.json")
        assert result is False

    def test_export_profile_writes_atomic_file(self, tmp_path: Path) -> None:
        """Hits the write+rename happy path including export_data fields."""
        exporter, pm = self._exporter()
        pm.get_profile = MagicMock(return_value=self._profile())
        out = tmp_path / "out.json"
        ok = exporter.export_profile("Test", out)
        assert ok is True
        assert out.exists()
        data = json.loads(out.read_text(encoding="utf-8"))
        assert data["profile_name"] == "Test"
        assert "exported_at" in data
        assert data["export_version"] == "1.0"
        # Temp file must be cleaned up.
        assert not (out.parent / f"{out.name}.tmp").exists()

    def test_preview_export_returns_none_when_profile_missing(self) -> None:
        """Hits the preview missing-profile branch (~line 263)."""
        exporter, pm = self._exporter()
        pm.get_profile = MagicMock(return_value=None)
        assert exporter.preview_export("ghost") is None

    def test_preview_export_calculates_statistics(self) -> None:
        """Hits lines 288-290 (statistics calculation)."""
        exporter, pm = self._exporter()
        pm.get_profile = MagicMock(return_value=self._profile())
        preview = exporter.preview_export("Test")
        assert preview is not None
        stats = preview["statistics"]
        assert stats["global_preferences_count"] == 1
        assert stats["directory_specific_count"] == 1
        assert stats["learned_patterns_count"] == 1
        assert stats["confidence_data_count"] == 1

    def test_export_selective_with_naming_and_folders(self, tmp_path: Path) -> None:
        """Hits lines 113-178 (export_selective happy path with naming + folders)."""
        exporter, pm = self._exporter()
        prof = self._profile(
            preferences={
                "global": {
                    "naming_patterns": {"docs": "{title}_{date}"},
                    "folder_mappings": {"src": "dst"},
                },
                "directory_specific": {"d": {"k": "v"}},
            }
        )
        pm.get_profile = MagicMock(return_value=prof)
        out = tmp_path / "selective.json"
        ok = exporter.export_selective(
            "Test", out, ["global", "naming", "folders", "learned_patterns", "confidence_data"]
        )
        assert ok is True
        data = json.loads(out.read_text(encoding="utf-8"))
        assert data["export_type"] == "selective"
        assert data["preferences"]["global"]["naming_patterns"] == {"docs": "{title}_{date}"}
        assert data["preferences"]["global"]["folder_mappings"] == {"src": "dst"}
        assert "learned_patterns" in data
        assert "confidence_data" in data

    def test_export_selective_returns_false_when_profile_missing(self, tmp_path: Path) -> None:
        """Hits lines 116-118 (selective profile-missing branch)."""
        exporter, pm = self._exporter()
        pm.get_profile = MagicMock(return_value=None)
        ok = exporter.export_selective("ghost", tmp_path / "x.json", ["global"])
        assert ok is False

    def test_validate_export_rejects_missing_preferences(self, tmp_path: Path) -> None:
        """Hits lines 215-217 (full export missing 'preferences' key)."""
        exporter, _ = self._exporter()
        f = tmp_path / "bad.json"
        f.write_text(
            json.dumps(
                {
                    "profile_name": "X",
                    "profile_version": "1.0",
                    "exported_at": "2026-05-03T00:00:00Z",
                    "export_type": "full",
                }
            )
        )
        assert exporter.validate_export(f) is False

    def test_validate_export_rejects_non_dict_preferences(self, tmp_path: Path) -> None:
        """Hits lines 220-222 (non-dict preferences)."""
        exporter, _ = self._exporter()
        f = tmp_path / "bad2.json"
        f.write_text(
            json.dumps(
                {
                    "profile_name": "X",
                    "profile_version": "1.0",
                    "exported_at": "2026-05-03T00:00:00Z",
                    "export_type": "full",
                    "preferences": "not-a-dict",
                }
            )
        )
        assert exporter.validate_export(f) is False

    def test_validate_export_rejects_missing_global_or_directory_keys(self, tmp_path: Path) -> None:
        """Hits lines 225-227 (preferences dict missing required keys)."""
        exporter, _ = self._exporter()
        f = tmp_path / "bad3.json"
        f.write_text(
            json.dumps(
                {
                    "profile_name": "X",
                    "profile_version": "1.0",
                    "exported_at": "2026-05-03T00:00:00Z",
                    "export_type": "full",
                    "preferences": {"global": {}},
                }
            )
        )
        assert exporter.validate_export(f) is False

    def test_validate_export_rejects_selective_missing_included_prefs(self, tmp_path: Path) -> None:
        """Hits lines 231-233 (selective export missing 'included_preferences')."""
        exporter, _ = self._exporter()
        f = tmp_path / "bad4.json"
        f.write_text(
            json.dumps(
                {
                    "profile_name": "X",
                    "profile_version": "1.0",
                    "exported_at": "2026-05-03T00:00:00Z",
                    "export_type": "selective",
                }
            )
        )
        assert exporter.validate_export(f) is False

    def test_validate_export_rejects_invalid_timestamp(self, tmp_path: Path) -> None:
        """Hits lines 238-240 (invalid ISO timestamp)."""
        exporter, _ = self._exporter()
        f = tmp_path / "bad5.json"
        f.write_text(
            json.dumps(
                {
                    "profile_name": "X",
                    "profile_version": "1.0",
                    "exported_at": "not-a-date",
                    "export_type": "selective",
                    "included_preferences": ["global"],
                }
            )
        )
        assert exporter.validate_export(f) is False

    def test_validate_export_rejects_invalid_json(self, tmp_path: Path) -> None:
        """Hits lines 244-246 (JSONDecodeError handler)."""
        exporter, _ = self._exporter()
        f = tmp_path / "broken.json"
        f.write_text("{ not json")
        assert exporter.validate_export(f) is False

    def test_validate_export_rejects_missing_file(self, tmp_path: Path) -> None:
        """Hits the file-not-found branch in validate_export."""
        exporter, _ = self._exporter()
        assert exporter.validate_export(tmp_path / "ghost.json") is False

    def test_export_multiple_aggregates_results(self, tmp_path: Path) -> None:
        """Hits lines 327-336 (export_multiple loop)."""
        exporter, pm = self._exporter()
        good = self._profile(name="good")
        pm.get_profile = MagicMock(side_effect=lambda name: good if name == "good" else None)
        results = exporter.export_multiple(["good", "missing"], tmp_path)
        assert results == {"good": True, "missing": False}
        assert (tmp_path / "good.json").exists()


# ---------------------------------------------------------------------------
# updater/installer.py: 71% -> 75% (+4pp)
# Targets missing lines 85-90, 92-95, 100, 152, 191, 209-211, 261-267.
# ---------------------------------------------------------------------------


class TestInstallerLift:
    """Lift integration coverage on src/updater/installer.py."""

    def test_score_asset_macos_universal_preferred(self) -> None:
        """Hits lines 85-90 (macOS scoring branches). _score_asset expects already-lowercased input."""
        from updater import installer

        with patch.object(installer.platform, "system", return_value="Darwin"):
            score_universal = installer._score_asset("fo-universal.tar.gz")
            score_dmg = installer._score_asset("fo-arm64.dmg")
            score_zip = installer._score_asset("fo-x86_64.zip")
        # universal gets +3 (also -3 for tar.gz) -> 0; dmg gets -5; zip gets -3.
        assert score_universal > score_zip
        assert score_zip > score_dmg

    def test_score_asset_windows_prefers_exe_over_setup(self) -> None:
        """Hits lines 92-95 (Windows scoring branches)."""
        from updater import installer

        with patch.object(installer.platform, "system", return_value="Windows"):
            score_exe = installer._score_asset("fo-1.0.exe")
            score_setup = installer._score_asset("fo-setup-1.0.exe")
        # exe gets +3; "setup" subtracts 4 -> setup-exe net = 3 - 4 = -1.
        assert score_exe > score_setup

    def test_score_asset_linux_prefers_appimage_over_tarball(self) -> None:
        """Hits lines 97-100 (Linux scoring branches). _score_asset takes a name_lower argument so we pass already-lowercased strings."""
        from updater import installer

        with patch.object(installer.platform, "system", return_value="Linux"):
            score_appimage = installer._score_asset("fo-1.0.appimage")
            score_tarball = installer._score_asset("fo-1.0.tar.gz")
        # AppImage +5, tar.gz +2.
        assert score_appimage > score_tarball

    def test_is_checksum_file_true_for_known_extensions(self) -> None:
        """Hits line 69 (_is_checksum_file true returns)."""
        from updater.installer import _is_checksum_file

        for name in ("file.sha256", "file.md5", "file.asc", "file.sig"):
            assert _is_checksum_file(name) is True
        assert _is_checksum_file("file.tar.gz") is False

    def test_get_arch_hints_x86_64(self) -> None:
        """Hits lines 49-50 (x86_64/amd64 hint branch)."""
        from updater import installer

        with (
            patch.object(installer.platform, "machine", return_value="x86_64"),
            patch.object(installer.platform, "system", return_value="Linux"),
        ):
            hints = installer._get_arch_hints()
        assert "x86_64" in hints
        assert "amd64" in hints

    def test_get_arch_hints_arm64_on_macos_includes_universal(self) -> None:
        """Hits lines 51-55 (arm64 + macOS universal append)."""
        from updater import installer

        with (
            patch.object(installer.platform, "machine", return_value="arm64"),
            patch.object(installer.platform, "system", return_value="Darwin"),
        ):
            hints = installer._get_arch_hints()
        assert "arm64" in hints
        assert "aarch64" in hints
        assert "universal" in hints
