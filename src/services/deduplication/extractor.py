# pyre-ignore-all-errors
"""Document text extraction module.

Extracts text content from various document formats for semantic analysis.
Supports PDF, DOCX, TXT, RTF, ODT, and Markdown document formats.

The public ``extract_text(file_path)`` entry point opens the file via
:class:`utils.safedir.SafeDir` so a symlink swapped into the organize root
between the directory walk and the read is refused with
``SymlinkRejected`` rather than dereferenced. On Windows (where SafeDir's
POSIX-only ``dir_fd`` + ``O_NOFOLLOW`` primitives are not available) the
legacy path-based extraction is used instead.

The underlying extractor methods now each accept either ``fileobj=`` (the
SafeDir-friendly entry point) or a path (the legacy entry point).
"""

from __future__ import annotations

import logging
import os
import sys
import xml.etree.ElementTree as _stdlib_ET  # fallback if defusedxml absent
import zipfile
from pathlib import Path
from typing import BinaryIO

try:
    # defusedxml prevents entity-bomb (billion-laughs) attacks in ODT XML.
    # It is a base dependency; the stdlib fallback is a safety net only.
    import defusedxml.ElementTree as _ET
except ImportError:  # pragma: no cover
    _ET = _stdlib_ET  # type: ignore[assignment]

from utils.safedir import SafeDir, SymlinkRejected

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extracts text content from various document formats.

    Supported formats:
    - PDF (.pdf)
    - Word Documents (.docx)
    - Plain Text (.txt, .md)
    - Rich Text (.rtf)
    - OpenDocument (.odt)
    """

    def __init__(self) -> None:
        """Initialize the document extractor."""
        self.supported_extensions = {".pdf", ".docx", ".txt", ".rtf", ".odt", ".md"}
        self._check_dependencies()

    def extract_text(self, file_path: Path) -> str:
        """Extract text from a single document.

        Opens the file via ``SafeDir.open_for_reader`` so a symlink swapped
        into the organize root between the directory walk and this read is
        refused — closes the symlink-following surface in the dedup
        ingestion pipeline (#264). On Windows the path-based fallback is
        used because SafeDir is POSIX-only.

        Args:
            file_path: Path to document file

        Returns:
            Extracted text content, or ``""`` when extraction fails for
            *any* reason — including a refused symlink (``SymlinkRejected``),
            an underlying ``OSError`` during the read, a malformed file, or
            a missing optional library. The empty-string contract is
            preserved from the pre-SafeDir implementation so consumers
            (``document_dedup`` etc.) behave identically: a file that
            can't be read simply contributes no text to the dedup corpus.

            **Read failures are NOT raised** — callers cannot rely on
            ``OSError`` to detect read problems. Only the missing-file
            and unsupported-format pre-checks raise.

        Raises:
            OSError: If the file does not exist (pre-flight check before
                any SafeDir / library work).
            ValueError: If the file's extension is not in
                ``supported_extensions``.
        """
        if not file_path.exists():
            raise OSError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if not self.supports_format(file_path):
            raise ValueError(f"Unsupported format: {extension}")

        # Try the SafeDir path first; fall back to the legacy path-branch
        # only when SafeDir is unavailable (Windows) — never on real FS
        # errors, which must propagate so the caller sees them.
        if sys.platform != "win32":
            try:
                with SafeDir.open_root(file_path.parent) as safe_dir:
                    fd = safe_dir.open_for_reader(file_path.name)
                    # fdopen takes ownership only once it returns; close the
                    # bare fd explicitly if fdopen itself raises (e.g. on
                    # a closed dir_fd race), otherwise the fd would leak.
                    try:
                        fileobj = os.fdopen(fd, "rb", closefd=True)
                    except OSError:
                        os.close(fd)
                        raise
                    with fileobj:
                        return self._extract_from_fileobj(extension, fileobj, file_path.name)
            except SymlinkRejected as exc:
                logger.warning(
                    "Refused to read symlinked file %s: %s", file_path, exc, exc_info=True
                )
                return ""
            except NotImplementedError:
                # SafeDir's POSIX primitives unavailable; fall through to
                # the path-based extraction below.
                logger.debug("SafeDir unavailable; using legacy reader for %s", file_path.name)
            except (OSError, ValueError, ImportError) as e:
                logger.error("Error extracting text from %s: %s", file_path, e, exc_info=True)
                return ""

        # Legacy path-based fallback (Windows) — preserves the original
        # API contract.
        try:
            return self._extract_via_path(extension, file_path)
        except (OSError, ValueError, ImportError) as e:
            logger.error("Error extracting text from %s: %s", file_path, e, exc_info=True)
            return ""

    def _extract_from_fileobj(self, extension: str, fileobj: BinaryIO, label: str) -> str:
        """Dispatch extraction to the right ``_extract_X`` with a fileobj."""
        if extension == ".pdf":
            return self._extract_pdf(fileobj=fileobj, label=label)
        if extension == ".docx":
            return self._extract_docx(fileobj=fileobj, label=label)
        if extension in (".txt", ".md"):
            return self._extract_text(fileobj=fileobj, label=label)
        if extension == ".rtf":
            return self._extract_rtf(fileobj=fileobj, label=label)
        if extension == ".odt":
            return self._extract_odt(fileobj=fileobj, label=label)
        logger.warning("No extractor for %s, treating as text", extension)
        return self._extract_text(fileobj=fileobj, label=label)

    def _extract_via_path(self, extension: str, file_path: Path) -> str:
        """Dispatch extraction to the right ``_extract_X`` with a path.

        Legacy fallback used only when SafeDir is unavailable.
        """
        if extension == ".pdf":
            return self._extract_pdf(file_path=file_path)
        if extension == ".docx":
            return self._extract_docx(file_path=file_path)
        if extension in (".txt", ".md"):
            return self._extract_text(file_path=file_path)
        if extension == ".rtf":
            return self._extract_rtf(file_path=file_path)
        if extension == ".odt":
            return self._extract_odt(file_path=file_path)
        logger.warning("No extractor for %s, treating as text", extension)
        return self._extract_text(file_path=file_path)

    def extract_batch(self, file_paths: list[Path]) -> dict[Path, str]:
        """Extract text from multiple documents in batch.

        Args:
            file_paths: List of document paths

        Returns:
            Dictionary mapping file paths to extracted text
        """
        results = {}

        for file_path in file_paths:
            try:
                text = self.extract_text(file_path)
                results[file_path] = text
                logger.debug(f"Extracted {len(text)} chars from {file_path.name}")
            except (OSError, ValueError, ImportError) as e:
                logger.warning(f"Failed to extract {file_path}: {e}")
                results[file_path] = ""

        logger.info(f"Batch extracted text from {len(results)} documents")

        return results

    def supports_format(self, file_path: Path) -> bool:
        """Check if a file format is supported.

        Args:
            file_path: Path to file

        Returns:
            True if format is supported
        """
        return file_path.suffix.lower() in self.supported_extensions

    def get_supported_formats(self) -> list[str]:
        """Get list of supported file formats.

        Returns:
            List of supported extensions
        """
        return sorted(self.supported_extensions)

    def _extract_pdf(
        self,
        file_path: Path | None = None,
        *,
        fileobj: BinaryIO | None = None,
        label: str | None = None,
    ) -> str:
        """Extract text from PDF file.

        Either ``file_path`` (legacy) or ``fileobj`` (SafeDir-friendly) must
        be supplied. The ``label`` arg is used in log messages when only a
        fileobj is given.
        """
        try:
            import pypdf

            try:
                from pypdf.errors import PyPdfError
            except (ImportError, AttributeError):
                PyPdfError = Exception  # type: ignore[misc,assignment]  # fallback when pypdf.errors unavailable

            text_parts: list[str] = []
            display = label or (file_path.name if file_path is not None else "<fileobj>")

            if fileobj is not None:
                pdf_reader = pypdf.PdfReader(fileobj)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            else:
                if file_path is None:
                    raise TypeError("path-branch requires file_path")
                # The opt-out marker below covers BOTH the bare open and the
                # nearby ``pypdf.PdfReader(f)`` call (within
                # ``_MARKER_WINDOW_BELOW=6``). Both belong to the same legacy
                # fallback block — when SafeDir is unavailable we accept the
                # path-based read end-to-end.
                with open(
                    file_path, "rb"
                ) as f:  # safedir: ok — Windows / NotImplementedError fallback
                    pdf_reader = pypdf.PdfReader(f)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from PDF: {display}")

            return full_text

        except ImportError:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            return ""
        except (PyPdfError, OSError, ValueError, KeyError, IndexError) as e:
            display = label or (file_path.name if file_path is not None else "<fileobj>")
            logger.error(f"Error extracting PDF {display}: {e}")
            return ""

    def _extract_docx(
        self,
        file_path: Path | None = None,
        *,
        fileobj: BinaryIO | None = None,
        label: str | None = None,
    ) -> str:
        """Extract text from DOCX file."""
        try:
            import docx

            if fileobj is not None:
                doc = docx.Document(fileobj)
            else:
                if file_path is None:
                    raise TypeError("path-branch requires file_path")
                doc = docx.Document(str(file_path))

            text_parts = [paragraph.text for paragraph in doc.paragraphs]

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            full_text = "\n".join(text_parts)
            display = label or (file_path.name if file_path is not None else "<fileobj>")
            logger.debug(f"Extracted {len(full_text)} chars from DOCX: {display}")

            return full_text

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return ""
        except (OSError, ValueError, KeyError) as e:
            display = label or (file_path.name if file_path is not None else "<fileobj>")
            logger.error(f"Error extracting DOCX {display}: {e}")
            return ""

    @staticmethod
    def _decode_bytes(raw: bytes) -> str:
        """Decode bytes using a utf-8 → cp1252 → latin-1 fallback chain.

        latin-1 is tried last because it accepts all 256 byte values without
        error; placing it earlier would make cp1252 unreachable.
        """
        for encoding in ("utf-8", "cp1252", "latin-1"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="ignore")  # unreachable; latin-1 always succeeds

    def _extract_text(
        self,
        file_path: Path | None = None,
        *,
        fileobj: BinaryIO | None = None,
        label: str | None = None,
    ) -> str:
        """Extract text from plain text file."""
        display = label or (file_path.name if file_path is not None else "<fileobj>")

        if fileobj is not None:
            try:
                raw = fileobj.read()
            except OSError as e:
                logger.error(f"Error reading text file {display}: {e}")
                return ""
            text = self._decode_bytes(raw)
            logger.debug(f"Read {len(text)} chars from text file: {display}")
            return text

        if file_path is None:
            raise TypeError("path-branch requires file_path")
        try:
            for encoding in ("utf-8", "latin-1", "cp1252", "ascii"):
                try:
                    with open(
                        file_path, encoding=encoding
                    ) as f:  # safedir: ok — Windows / NotImplementedError fallback
                        text = f.read()
                    logger.debug(f"Read {len(text)} chars from text file: {display}")
                    return text
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, "rb") as f:  # safedir: ok — Windows / NotImplementedError fallback
                text = f.read().decode("utf-8", errors="ignore")

            return text

        except OSError as e:
            logger.error(f"Error reading text file {display}: {e}")
            return ""

    def _extract_rtf(
        self,
        file_path: Path | None = None,
        *,
        fileobj: BinaryIO | None = None,
        label: str | None = None,
    ) -> str:
        """Extract text from RTF file."""
        display = label or (file_path.name if file_path is not None else "<fileobj>")

        try:
            if fileobj is not None:
                rtf_content = self._decode_bytes(fileobj.read())
            else:
                if file_path is None:
                    raise TypeError("path-branch requires file_path")
                with open(
                    file_path, "rb"
                ) as f:  # safedir: ok — Windows / NotImplementedError fallback
                    rtf_content = self._decode_bytes(f.read())

            # Try using striprtf if available
            try:
                from striprtf.striprtf import rtf_to_text

                text = str(rtf_to_text(rtf_content))
                logger.debug(f"Extracted {len(text)} chars from RTF: {display}")
                return text

            except ImportError:
                # Fallback: simple RTF stripping
                logger.warning("striprtf not installed, using basic extraction")
                import re

                text = re.sub(r"\\[a-z]+\d*\s?", "", rtf_content)
                text = re.sub(r"[{}]", "", text)
                text = text.strip()
                return text

        except (OSError, ValueError, ImportError) as e:
            logger.error(f"Error extracting RTF {display}: {e}")
            return ""

    def _extract_odt(
        self,
        file_path: Path | None = None,
        *,
        fileobj: BinaryIO | None = None,
        label: str | None = None,
    ) -> str:
        """Extract text from ODT file."""
        display = label or (file_path.name if file_path is not None else "<fileobj>")

        try:
            # ODT files are ZIP archives — zipfile.ZipFile accepts both
            # paths and file-like objects.
            source = fileobj if fileobj is not None else file_path
            if source is None:
                raise TypeError("_extract_odt requires file_path or fileobj")
            with zipfile.ZipFile(source, "r") as odt_zip:
                content_xml = odt_zip.read("content.xml")

            root = _ET.fromstring(content_xml)

            namespaces = {
                "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
                "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
            }

            text_parts = []

            for paragraph in root.findall(".//text:p", namespaces):
                if paragraph.text:
                    text_parts.append(paragraph.text)

                for child in paragraph:
                    if child.text:
                        text_parts.append(child.text)
                    if child.tail:
                        text_parts.append(child.tail)

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from ODT: {display}")

            return full_text

        except (OSError, KeyError, ValueError, zipfile.BadZipFile) as e:
            logger.error(f"Error extracting ODT {display}: {e}")
            return ""

    def _check_dependencies(self) -> None:
        """Check if required dependencies are installed."""
        dependencies = {
            "pypdf": "PDF extraction",
            "docx": "DOCX extraction",
        }

        missing = []
        for module, purpose in dependencies.items():
            try:
                __import__(module if module != "docx" else "docx")
            except ImportError:
                missing.append(f"{module} ({purpose})")

        if missing:
            logger.warning(
                f"Missing optional dependencies: {', '.join(missing)}. "
                "Some document formats may not be supported."
            )
