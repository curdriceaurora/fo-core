"""Document text extraction module.

Extracts text content from various document formats for semantic analysis.
Supports PDF, DOCX, TXT, RTF, ODT, and Markdown document formats.
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extracts text content from various document formats.

    Supported formats:
    - PDF (.pdf)
    - Word Documents (.docx)
    - Plain Text (.txt, .md)
    - Rich Text (.rtf)
    - OpenDocument (.odt)
    """

    def __init__(self) -> None:
        """Initialize the document extractor."""
        self.supported_extensions = {".pdf", ".docx", ".txt", ".rtf", ".odt", ".md"}
        self._check_dependencies()

    def extract_text(self, file_path: Path) -> str:
        """Extract text from a single document.

        Args:
            file_path: Path to document file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format not supported
            OSError: If file cannot be read
        """
        if not file_path.exists():
            raise OSError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if not self.supports_format(file_path):
            raise ValueError(f"Unsupported format: {extension}")

        try:
            if extension == ".pdf":
                return self._extract_pdf(file_path)
            elif extension == ".docx":
                return self._extract_docx(file_path)
            elif extension == ".txt" or extension == ".md":
                return self._extract_text(file_path)
            elif extension == ".rtf":
                return self._extract_rtf(file_path)
            elif extension == ".odt":
                return self._extract_odt(file_path)
            else:
                logger.warning(f"No extractor for {extension}, treating as text")
                return self._extract_text(file_path)

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def extract_batch(self, file_paths: list[Path]) -> dict[Path, str]:
        """Extract text from multiple documents in batch.

        Args:
            file_paths: List of document paths

        Returns:
            Dictionary mapping file paths to extracted text
        """
        results = {}

        for file_path in file_paths:
            try:
                text = self.extract_text(file_path)
                results[file_path] = text
                logger.debug(f"Extracted {len(text)} chars from {file_path.name}")
            except Exception as e:
                logger.warning(f"Failed to extract {file_path}: {e}")
                results[file_path] = ""

        logger.info(f"Batch extracted text from {len(results)} documents")

        return results

    def supports_format(self, file_path: Path) -> bool:
        """Check if a file format is supported.

        Args:
            file_path: Path to file

        Returns:
            True if format is supported
        """
        return file_path.suffix.lower() in self.supported_extensions

    def get_supported_formats(self) -> list[str]:
        """Get list of supported file formats.

        Returns:
            List of supported extensions
        """
        return sorted(self.supported_extensions)

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text
        """
        try:
            import pypdf

            text_parts = []

            with open(file_path, "rb") as f:
                pdf_reader = pypdf.PdfReader(f)

                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from PDF: {file_path.name}")

            return full_text

        except ImportError:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            return ""
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            return ""

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text
        """
        try:
            import docx

            doc = docx.Document(str(file_path))

            # Extract text from all paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from DOCX: {file_path.name}")

            return full_text

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            return ""

    def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text file.

        Args:
            file_path: Path to text file

        Returns:
            File contents
        """
        try:
            # Try multiple encodings
            encodings = ["utf-8", "latin-1", "cp1252", "ascii"]

            for encoding in encodings:
                try:
                    with open(file_path, encoding=encoding) as f:
                        text = f.read()
                    logger.debug(f"Read {len(text)} chars from text file: {file_path.name}")
                    return text
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, "rb") as f:
                text = f.read().decode("utf-8", errors="ignore")

            return text

        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""

    def _extract_rtf(self, file_path: Path) -> str:
        """Extract text from RTF file.

        Args:
            file_path: Path to RTF file

        Returns:
            Extracted text
        """
        try:
            # Try using striprtf if available
            try:
                from striprtf.striprtf import rtf_to_text

                with open(file_path, encoding="utf-8", errors="ignore") as f:
                    rtf_content = f.read()

                text = str(rtf_to_text(rtf_content))
                logger.debug(f"Extracted {len(text)} chars from RTF: {file_path.name}")

                return text

            except ImportError:
                # Fallback: simple RTF stripping
                logger.warning("striprtf not installed, using basic extraction")

                with open(file_path, encoding="utf-8", errors="ignore") as f:
                    content = f.read()

                # Very basic RTF stripping (removes control words)
                import re

                text = re.sub(r"\\[a-z]+\d*\s?", "", content)
                text = re.sub(r"[{}]", "", text)
                text = text.strip()

                return text

        except Exception as e:
            logger.error(f"Error extracting RTF {file_path}: {e}")
            return ""

    def _extract_odt(self, file_path: Path) -> str:
        """Extract text from ODT file.

        Args:
            file_path: Path to ODT file

        Returns:
            Extracted text
        """
        try:
            import xml.etree.ElementTree as ET
            import zipfile

            # ODT files are ZIP archives
            with zipfile.ZipFile(file_path, "r") as odt_zip:
                # Extract content.xml
                content_xml = odt_zip.read("content.xml")

            # Parse XML
            root = ET.fromstring(content_xml)

            # Extract all text nodes
            # ODT uses OpenDocument namespace
            namespaces = {
                "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
                "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
            }

            text_parts = []

            # Find all text:p (paragraph) elements
            for paragraph in root.findall(".//text:p", namespaces):
                if paragraph.text:
                    text_parts.append(paragraph.text)

                # Also get text from child elements
                for child in paragraph:
                    if child.text:
                        text_parts.append(child.text)
                    if child.tail:
                        text_parts.append(child.tail)

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from ODT: {file_path.name}")

            return full_text

        except Exception as e:
            logger.error(f"Error extracting ODT {file_path}: {e}")
            return ""

    def _check_dependencies(self) -> None:
        """Check if required dependencies are installed."""
        dependencies = {
            "pypdf": "PDF extraction",
            "docx": "DOCX extraction",
        }

        missing = []
        for module, purpose in dependencies.items():
            try:
                __import__(module if module != "docx" else "docx")
            except ImportError:
                missing.append(f"{module} ({purpose})")

        if missing:
            logger.warning(
                f"Missing optional dependencies: {', '.join(missing)}. "
                "Some document formats may not be supported."
            )
