#!/usr/bin/env python3
"""Generate real sample files for the E2E test fixture.

Run once to populate tests/fixtures/e2e_samples/:
    python scripts/generate_e2e_samples.py

Requires: python-docx, openpyxl, Pillow (all already in dev deps)
Output:   5 tiny valid files totalling < 100 KB committed to the repo
"""

from __future__ import annotations

import wave
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent.parent / "tests" / "fixtures" / "e2e_samples"


def generate_docx(path: Path) -> None:
    """Create a minimal valid .docx via python-docx."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise SystemExit("python-docx not installed: pip install python-docx") from exc

    # NOTE: Document() loads python-docx's bundled default.docx template (~36 KB).
    # This dominates the total fixture size. Use a stripped template if the 100 KB
    # budget ever tightens (see: python-docx docs on custom templates).
    doc = Document()
    doc.add_heading("E2E Sample Document", level=1)
    doc.add_paragraph("This is a minimal Word document used as a fixture in the E2E test suite.")
    doc.add_paragraph("It contains enough structure to be a valid DOCX file.")
    doc.save(path)


def generate_xlsx(path: Path) -> None:
    """Create a minimal valid .xlsx via openpyxl."""
    try:
        import openpyxl  # type: ignore[import-untyped]
    except ImportError as exc:
        raise SystemExit("openpyxl not installed: pip install openpyxl") from exc

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sample"  # type: ignore[union-attr]
    ws.append(["Name", "Value", "Date"])  # type: ignore[union-attr]
    ws.append(["Item A", 100, "2024-01-01"])  # type: ignore[union-attr]
    ws.append(["Item B", 200, "2024-01-02"])  # type: ignore[union-attr]
    wb.save(path)


def _generate_image(path: Path, fmt: str) -> None:
    """Create a minimal valid 1×1 red pixel image via Pillow."""
    try:
        from PIL import Image  # type: ignore[import-untyped]
    except ImportError as exc:
        raise SystemExit("Pillow not installed: pip install Pillow") from exc

    img = Image.new("RGB", (1, 1), color=(255, 0, 0))
    img.save(path, format=fmt)


def generate_jpg(path: Path) -> None:
    """Create a minimal valid 1×1 red pixel JPEG via Pillow.

    NOTE: JPEG is lossy — pixel (255,0,0) round-trips as (254,0,0).
    Tests must not assert exact pixel values from this file.
    """
    _generate_image(path, "JPEG")


def generate_png(path: Path) -> None:
    """Create a minimal valid 1×1 red pixel PNG via Pillow."""
    _generate_image(path, "PNG")


def generate_wav(path: Path) -> None:
    """Create a minimal valid WAV file (0.1s silence, 16-bit mono 8kHz)."""
    sample_rate = 8000
    num_samples = int(sample_rate * 0.1)  # 0.1 seconds of silence
    with wave.open(str(path), "w") as wf:
        wf.setnchannels(1)
        wf.setsampwidth(2)  # 16-bit
        wf.setframerate(sample_rate)
        # bytes(n) produces n zero bytes; 2 bytes per 16-bit sample = num_samples frames
        wf.writeframes(bytes(num_samples * 2))


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    generators = [
        ("sample.docx", generate_docx),
        ("sample.xlsx", generate_xlsx),
        ("sample.jpg", generate_jpg),
        ("sample.png", generate_png),
        ("sample.wav", generate_wav),
    ]

    total_bytes = 0
    for filename, generator in generators:
        output_path = OUTPUT_DIR / filename
        if output_path.exists():
            size = output_path.stat().st_size
            total_bytes += size
            print(f"  {filename}: {size:,} bytes (skipped, already exists)")
        else:
            generator(output_path)
            size = output_path.stat().st_size
            total_bytes += size
            print(f"  {filename}: {size:,} bytes")

    print(f"\nTotal: {total_bytes:,} bytes ({total_bytes / 1024:.1f} KB)")
    print(f"Output: {OUTPUT_DIR}")

    if total_bytes > 100 * 1024:
        print(f"\nWARNING: Total size {total_bytes / 1024:.1f} KB exceeds 100 KB target.")


if __name__ == "__main__":
    main()
